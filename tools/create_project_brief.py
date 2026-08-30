from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "DiskWise_Project_Brief.docx"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None: node = OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.first_child_found_in('w:tcW')
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa'); set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

doc = Document(); sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(38,56,74)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in [('Heading 1',16,18,10,'2E74B5'),('Heading 2',13,14,7,'2E74B5'),('Heading 3',12,10,5,'1F4D78')]:
    s = doc.styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)

title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after = Pt(3)
r = title.add_run('磁盘智理 / DiskWise'); r.bold=True; r.font.name='Calibri'; r.font.size=Pt(26); r.font.color.rgb=RGBColor(11,37,69)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.add_run('Project handoff and delivery brief').italic = True
doc.add_paragraph('文档用途：为后续 Agent、维护者和 GitHub 发布准备提供统一的项目背景、当前成果、运行方式与交付注意事项。')

doc.add_heading('1. 项目概况', level=1)
doc.add_paragraph('DiskWise 是一款 Windows 磁盘空间分析与安全文件处置工具。它支持目录导航、文件关联软件识别、大文件与大文件夹扫描、垃圾/临时文件识别、删除影响提示，以及回收站和永久删除操作。')

doc.add_heading('2. 当前成果', level=1)
table = doc.add_table(rows=1, cols=3); table.alignment=WD_TABLE_ALIGNMENT.CENTER; set_table_widths(table,[2200,3000,4160])
for c, text in zip(table.rows[0].cells,['能力','状态','说明']): c.text=text; shade(c,'E8EEF5')
rows=[('应用启动','已完成','正式入口为 src/main.py；动态配置 PyQt5 插件路径，支持中文安装路径。'),('空间分析','已完成','大文件、大文件夹、实际占用、逻辑大小和垃圾/临时文件扫描。'),('安全处置','已完成','删除建议、云盘风险提示、移入回收站、永久删除、右键菜单。'),('界面体验','已完成','中文/英文切换、长路径省略、文件夹递归容量、三角形数值箭头。'),('发布资源','已完成','assets Logo、PyInstaller spec、Inno Setup 脚本和便携版 ZIP。')]
for row in rows:
    cells=table.add_row().cells
    for c,text in zip(cells,row): c.text=text

doc.add_heading('3. 目录结构（英文命名）', level=1)
for text in ['src/：源代码、正式入口 main.py 和回归测试。','assets/：窗口、任务栏、快捷方式和安装包使用的 Logo。','packaging/：DiskWise.spec、DiskWise.iss 与打包说明。','tools/：仅用于维护和截图的辅助脚本。','docs/archive/：历史交接与阶段记录，不参与运行。']:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('4. 运行与验证', level=1)
doc.add_paragraph('开发环境运行命令：')
doc.add_paragraph('& ".\\.venv\\Scripts\\python.exe" -X utf8 ".\\src\\main.py"', style='Intense Quote')
doc.add_paragraph('回归测试：')
doc.add_paragraph('& ".\\.venv\\Scripts\\python.exe" -X utf8 ".\\src\\test_phase3.py"', style='Intense Quote')
doc.add_paragraph('已验证内容包括：PyQt5 插件初始化、有效 QStyle 图标、C/D 磁盘导航、历史返回、文件详情、关联软件识别、扫描结果右键菜单、垃圾文件清理和英文界面。')

doc.add_heading('5. 打包与发布', level=1)
doc.add_paragraph('PyInstaller 配置已准备好。中文路径下构建时需使用临时盘符映射规避 PyInstaller Qt 插件解析限制，映射完成后必须解除；该映射不是应用功能，也不会成为项目盘符。当前已生成：')
doc.add_paragraph('installer-output/DiskWise-Portable-1.0.0.zip', style='List Bullet')
doc.add_paragraph('真正的 Setup.exe 需要在安装 Inno Setup 后编译 packaging/DiskWise.iss。')

doc.add_heading('6. 后续 Agent 注意事项', level=1)
for text in ['不要恢复中文源文件名；代码导入和打包配置均使用英文模块名。','正式启动只能使用 src/main.py，不要直接运行 main_window.py。','不要把 .venv、build、dist、installer-output 或缓存提交到 GitHub。','删除操作必须保留确认流程；云盘、系统目录和程序目录要继续显示风险提示。','如修改界面或扫描逻辑，先运行语法检查和 test_phase3.py，再启动正式入口做视觉验证。']:
    doc.add_paragraph(text, style='List Bullet')

doc.add_heading('7. 已知限制', level=1)
doc.add_paragraph('当前机器未安装 Inno Setup，因此仓库内提供完整 ISS 配置和已验证的便携版 ZIP，但未生成单文件安装器。PyInstaller 构建日志中的 NumPy hook 警告来自混用的全局 PyInstaller 环境，不影响本项目构建结果；后续可在干净的 Python/打包环境中消除该警告。')

footer = sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run('DiskWise project brief')
doc.save(OUT); print(OUT)
